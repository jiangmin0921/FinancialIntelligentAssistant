"""
自动生成企业财务文档（PDF/Word格式）
"""

import os
from datetime import datetime
from pathlib import Path
from typing import List

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class DocumentGenerator:
    """文档生成器"""
    # 文档生成器用于自动生成企业级的财务制度文档、报销流程、财务问答样例等内容，
    # 可输出 PDF 或 Word 文件，便于测试和数据准备。
    # 支持中文显示，自动注册常见字体，仅依赖于 reportlab 和 python-docx。
    #
    # 通常调用 generate_financial_documents() 方法生成固定模板或样例内容。
    def __init__(self, output_dir: str = "./data/generated"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # 注册中文字体
        if HAS_REPORTLAB:
            self._register_chinese_fonts()
    
    def _register_chinese_fonts(self):
        """注册中文字体"""
        try:
            import platform
            
            font_registered = False
            
            # Windows系统字体路径
            if platform.system() == 'Windows':
                font_paths = [
                    (r'C:\Windows\Fonts\simsun.ttc', 0, 'ChineseFont'),      # 宋体
                    (r'C:\Windows\Fonts\simhei.ttf', None, 'ChineseFont'),  # 黑体
                    (r'C:\Windows\Fonts\msyh.ttc', 0, 'ChineseFont'),      # 微软雅黑
                    (r'C:\Windows\Fonts\msyhbd.ttc', 0, 'ChineseFontBold'), # 微软雅黑 Bold
                ]
                
                # 尝试注册字体
                for font_path, subfont_index, font_name in font_paths:
                    if os.path.exists(font_path):
                        try:
                            if subfont_index is not None and font_path.endswith('.ttc'):
                                # TTC文件需要指定字体索引
                                pdfmetrics.registerFont(TTFont(font_name, font_path, subfontIndex=subfont_index))
                            else:
                                pdfmetrics.registerFont(TTFont(font_name, font_path))
                            
                            font_registered = True
                            print(f"[OK] 已注册中文字体: {os.path.basename(font_path)}")
                            
                            # 如果没有粗体字体，使用同一个字体作为粗体
                            if font_name == 'ChineseFont' and 'ChineseFontBold' not in pdfmetrics.getRegisteredFontNames():
                                try:
                                    pdfmetrics.registerFont(TTFont('ChineseFontBold', font_path, subfontIndex=subfont_index if subfont_index is not None else 0))
                                except:
                                    pass
                            
                            break
                        except Exception as e:
                            continue
                
                if not font_registered:
                    print("[WARNING] 未找到系统中文字体")
            else:
                # Linux/Mac系统 - 尝试常见字体路径
                linux_font_paths = [
                    '/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc',
                    '/System/Library/Fonts/STHeiti Light.ttc',  # Mac
                ]
                
                for font_path in linux_font_paths:
                    if os.path.exists(font_path):
                        try:
                            pdfmetrics.registerFont(TTFont('ChineseFont', font_path, subfontIndex=0))
                            font_registered = True
                            print(f"[OK] 已注册中文字体: {font_path}")
                            break
                        except:
                            continue
                
                if not font_registered:
                    print("[WARNING] 未找到中文字体，请安装中文字体包")
            
            # 检查是否注册成功
            if 'ChineseFont' not in pdfmetrics.getRegisteredFontNames():
                print("[WARNING] 中文字体注册失败，PDF可能显示乱码")
                print("  建议：安装中文字体或使用Word文档格式")
                
        except Exception as e:
            print(f"[ERROR] 注册中文字体时出错: {e}")
            print("将使用默认字体（可能不支持中文）")
        
    def generate_financial_documents(self) -> List[str]:
        """生成所有财务文档"""
        generated_files = []
        
        # 生成PDF文档
        if HAS_REPORTLAB:
            pdf_files = self._generate_pdf_documents()
            generated_files.extend(pdf_files)
        else:
            print("警告: reportlab未安装，跳过PDF生成")
        
        # 生成Word文档
        if HAS_DOCX:
            word_files = self._generate_word_documents()
            generated_files.extend(word_files)
        else:
            print("警告: python-docx未安装，跳过Word生成")
        
        return generated_files
    
    def _generate_pdf_documents(self) -> List[str]:
        """生成PDF文档"""
        pdf_files = []
        
        # 1. 差旅费报销制度
        content = {
            "title": "差旅费报销制度 V3.2",
            "sections": [
                {
                    "title": "第一条 适用范围",
                    "content": "本制度适用于公司全体员工因公出差产生的差旅费用报销。包括但不限于：交通费、住宿费、餐费、通讯费等。"
                },
                {
                    "title": "第二条 交通费标准",
                    "content": "1. 国内机票：报销上限为经济舱全价票的80%，需提供行程单和登机牌作为凭证。\n2. 高铁/动车：可报销二等座及以下标准。\n3. 市内交通：出租车需提供发票，单次超过100元需说明原因。"
                },
                {
                    "title": "第三条 住宿费标准",
                    "content": "1. 一线城市（北上广深）：住宿标准不超过500元/晚。\n2. 二线城市：住宿标准不超过350元/晚。\n3. 其他城市：住宿标准不超过250元/晚。\n4. 需提供酒店发票和住宿明细。"
                },
                {
                    "title": "第四条 餐费标准",
                    "content": "1. 早餐：不超过50元/人。\n2. 午餐：不超过100元/人。\n3. 晚餐：不超过150元/人。\n4. 需提供餐饮发票，超过标准部分需说明原因。"
                },
                {
                    "title": "第五条 报销流程",
                    "content": "1. 出差前需在OA系统提交出差申请，获得审批。\n2. 出差结束后7个工作日内提交报销申请。\n3. 报销周期：每月1-10日受理上月单据。\n4. 单张发票超过3000元需主管审批。\n5. 所有票据需真实有效，不得涂改。"
                }
            ]
        }
        pdf_file = self._create_pdf("差旅费报销制度_V3.2.pdf", content)
        pdf_files.append(pdf_file)
        
        # 2. 财务管理制度
        content = {
            "title": "财务管理制度",
            "sections": [
                {
                    "title": "第一章 总则",
                    "content": "为规范公司财务管理，提高资金使用效率，保障公司资产安全，根据国家相关法律法规，结合公司实际情况，制定本制度。"
                },
                {
                    "title": "第二章 预算管理",
                    "content": "1. 公司实行全面预算管理，各部门需在每年12月提交下年度预算。\n2. 预算执行过程中，如发生重大变化，需及时调整预算。\n3. 财务部门负责预算执行情况的监督和考核。"
                },
                {
                    "title": "第三章 费用管理",
                    "content": "1. 所有费用支出需符合公司费用标准。\n2. 费用报销需提供真实有效的票据。\n3. 大额费用（超过5000元）需提前申请审批。"
                }
            ]
        }
        pdf_file = self._create_pdf("财务管理制度.pdf", content)
        pdf_files.append(pdf_file)
        
        # 3. 财务FAQ
        content = {
            "title": "财务常见问题解答（FAQ）",
            "sections": [
                {
                    "title": "Q1: 报销需要哪些材料？",
                    "content": "A: 报销需要提供：1) 费用发票原件；2) 费用明细清单；3) 相关审批单据；4) 出差申请单（如为差旅费）。"
                },
                {
                    "title": "Q2: 报销周期是多久？",
                    "content": "A: 报销周期为每月1-10日受理上月单据，财务部门在收到完整材料后5个工作日内完成审核和付款。"
                },
                {
                    "title": "Q3: 发票丢失怎么办？",
                    "content": "A: 发票丢失需提供发票复印件或相关证明材料，并填写《发票丢失说明表》，经部门主管审批后可报销。"
                },
                {
                    "title": "Q4: 可以报销哪些费用？",
                    "content": "A: 可以报销的费用包括：差旅费、交通费、业务招待费、办公用品费、通讯费等因公产生的合理费用。"
                }
            ]
        }
        pdf_file = self._create_pdf("财务FAQ.pdf", content)
        pdf_files.append(pdf_file)
        
        return pdf_files
    
    def _create_pdf(self, filename: str, content: dict) -> str:
        """创建PDF文档"""
        filepath = self.output_dir / filename
        
        # 如果文件已存在，先删除
        if filepath.exists():
            try:
                filepath.unlink()
            except PermissionError:
                # 如果文件被占用，尝试重命名
                import time
                backup_name = filepath.stem + "_backup_" + str(int(time.time())) + filepath.suffix
                backup_path = filepath.parent / backup_name
                try:
                    filepath.rename(backup_path)
                except:
                    raise PermissionError(f"无法删除或重命名文件 {filename}，文件可能正在被其他程序打开")
            except Exception as e:
                raise Exception(f"处理已存在文件时出错: {e}")
        
        doc = SimpleDocTemplate(str(filepath), pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # 确定使用的中文字体
        chinese_font = 'ChineseFont' if 'ChineseFont' in pdfmetrics.getRegisteredFontNames() else 'Helvetica'
        chinese_font_bold = 'ChineseFontBold' if 'ChineseFontBold' in pdfmetrics.getRegisteredFontNames() else chinese_font
        
        # 标题样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=chinese_font_bold,
            fontSize=18,
            textColor='#2563eb',
            spaceAfter=30,
            alignment=1  # 居中
        )
        
        # 章节标题样式
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontName=chinese_font_bold,
            fontSize=14,
            textColor='#1f2937',
            spaceAfter=12,
            spaceBefore=20
        )
        
        # 正文样式
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=chinese_font,
            fontSize=11,
            leading=16,
            spaceAfter=12
        )
        
        # 添加标题
        story.append(Paragraph(content["title"], title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # 添加章节
        for section in content["sections"]:
            story.append(Paragraph(section["title"], heading_style))
            # 处理换行
            paragraphs = section["content"].split('\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), normal_style))
            story.append(Spacer(1, 0.2*inch))
        
        doc.build(story)
        return str(filepath)
    
    def _generate_word_documents(self) -> List[str]:
        """生成Word文档"""
        word_files = []
        
        # 1. 费用报销流程指南
        content = {
            "title": "费用报销流程指南",
            "sections": [
                {
                    "title": "一、报销前准备",
                    "content": [
                        "1. 确认费用符合公司报销标准",
                        "2. 收集并整理所有相关票据",
                        "3. 填写费用报销单",
                        "4. 准备相关审批材料"
                    ]
                },
                {
                    "title": "二、提交报销",
                    "content": [
                        "1. 登录OA系统，进入报销模块",
                        "2. 选择报销类型（差旅费/交通费/业务招待费等）",
                        "3. 上传发票和凭证照片",
                        "4. 填写费用明细和说明",
                        "5. 提交审批"
                    ]
                },
                {
                    "title": "三、审批流程",
                    "content": [
                        "1. 部门主管初审（1-2个工作日）",
                        "2. 财务部门审核（2-3个工作日）",
                        "3. 超过5000元需总经理审批",
                        "4. 审批通过后，财务部门安排付款"
                    ]
                }
            ]
        }
        word_file = self._create_word("费用报销流程指南.docx", content)
        word_files.append(word_file)
        
        # 2. 财务制度补充说明
        content = {
            "title": "财务制度补充说明",
            "sections": [
                {
                    "title": "发票要求",
                    "content": [
                        "1. 发票必须真实有效，不得涂改",
                        "2. 发票抬头必须为公司全称",
                        "3. 发票内容需与实际业务相符",
                        "4. 电子发票需打印并签字确认"
                    ]
                },
                {
                    "title": "报销时限",
                    "content": [
                        "1. 差旅费用需在返回后7个工作日内报销",
                        "2. 其他费用需在发生当月报销",
                        "3. 超过3个月的票据不予报销",
                        "4. 特殊情况需提前申请延期"
                    ]
                },
                {
                    "title": "注意事项",
                    "content": [
                        "1. 报销前请仔细核对票据金额",
                        "2. 确保所有信息填写完整准确",
                        "3. 如有疑问，及时联系财务部门",
                        "4. 保留报销凭证复印件备查"
                    ]
                }
            ]
        }
        word_file = self._create_word("财务制度补充说明.docx", content)
        word_files.append(word_file)
        
        return word_files
    
    def _create_word(self, filename: str, content: dict) -> str:
        """创建Word文档"""
        filepath = self.output_dir / filename
        
        # 如果文件已存在，先删除
        if filepath.exists():
            try:
                filepath.unlink()
            except PermissionError:
                # 如果文件被占用，尝试重命名
                import time
                backup_name = filepath.stem + "_backup_" + str(int(time.time())) + filepath.suffix
                backup_path = filepath.parent / backup_name
                try:
                    filepath.rename(backup_path)
                except:
                    raise PermissionError(f"无法删除或重命名文件 {filename}，文件可能正在被其他程序打开")
            except Exception as e:
                raise Exception(f"处理已存在文件时出错: {e}")
        
        doc = Document()
        
        # 添加标题
        title = doc.add_heading(content["title"], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加日期
        date_para = doc.add_paragraph(f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # 空行
        
        # 添加章节
        for section in content["sections"]:
            # 章节标题
            heading = doc.add_heading(section["title"], level=1)
            
            # 章节内容
            for item in section["content"]:
                para = doc.add_paragraph(item, style='List Bullet')
                para.paragraph_format.space_after = Pt(6)
            
            doc.add_paragraph()  # 空行
        
        doc.save(str(filepath))
        return str(filepath)


if __name__ == "__main__":
    generator = DocumentGenerator()
    files = generator.generate_financial_documents()
    print(f"已生成 {len(files)} 个文档：")
    for file in files:
        print(f"  - {file}")

